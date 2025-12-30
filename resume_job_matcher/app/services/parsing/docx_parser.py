from __future__ import annotations

import io

from docx import Document


class DocxParseError(RuntimeError):
    pass


def extract_text_from_docx_bytes(docx_bytes: bytes) -> str:
    """Extract text from DOCX bytes via python-docx (paragraphs + tables)."""
    try:
        doc = Document(io.BytesIO(docx_bytes))
    except Exception as e:
        raise DocxParseError(f"Failed to read DOCX: {e}") from e

    parts: list[str] = []

    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)

    # Tables are common in resumes; include cell text too.
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join((cell.text or "").strip() for cell in row.cells).strip()
            if row_text and row_text != "|":
                parts.append(row_text)

    text = "\n".join(parts).strip()
    if not text:
        raise DocxParseError("No extractable text found in DOCX.")
    return text






