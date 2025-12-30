from __future__ import annotations

import io
from dataclasses import dataclass
from typing import Any, Literal

from docx import Document


@dataclass(frozen=True)
class DocxExportResult:
    filename: str
    content_type: str
    data: bytes


def _safe_filename(name: str) -> str:
    return "".join(c for c in name if c.isalnum() or c in ("-", "_", ".", " ")).strip().replace(" ", "_")


def export_docx(*, analysis: dict[str, Any], mode: Literal["resume_bullets", "cover_letter"]) -> DocxExportResult:
    analysis_id = analysis.get("analysis_id") or "analysis"
    score = (analysis.get("score") or {}).get("final_match_score")
    skill_gap = analysis.get("skill_gap") or {}
    suggestions = analysis.get("suggestions") or {}

    doc = Document()

    if mode == "resume_bullets":
        doc.add_heading("Tailored Resume Bullets", level=1)
        doc.add_paragraph(f"Analysis: {analysis_id}")
        if score is not None:
            doc.add_paragraph(f"Match score: {round(float(score), 2)} / 100")

        doc.add_heading("Key Strengths", level=2)
        for s in (suggestions.get("key_strengths") or [])[:8]:
            doc.add_paragraph(str(s), style="List Bullet")

        doc.add_heading("Missing Required Skills to Add", level=2)
        missing = (skill_gap.get("missing_required_skills") or [])[:12]
        if not missing:
            doc.add_paragraph("None detected.", style=None)
        else:
            for s in missing:
                doc.add_paragraph(str(s), style="List Bullet")

        doc.add_heading("Bullet Rewrites (Before → After)", level=2)
        rewrites = (suggestions.get("bullet_rewrites") or [])[:6]
        if not rewrites:
            doc.add_paragraph("No bullet rewrites available for this analysis.")
        else:
            for b in rewrites:
                before = (b or {}).get("before") or ""
                after = (b or {}).get("after") or ""
                if before:
                    doc.add_paragraph("Before: " + before)
                if after:
                    doc.add_paragraph("After: " + after)
                doc.add_paragraph("")  # spacer

        filename = _safe_filename(f"resume_bullets_{analysis_id}.docx")

    else:
        doc.add_heading("Cover Letter Draft", level=1)
        doc.add_paragraph(f"Analysis: {analysis_id}")

        missing = (skill_gap.get("missing_required_skills") or [])[:6]
        matching = (skill_gap.get("matching_skills") or [])[:10]

        doc.add_paragraph("Dear Hiring Manager,")
        doc.add_paragraph(
            "I’m excited to apply for this role. Based on the job requirements and my background, "
            "I bring strong alignment in key areas while actively strengthening the remaining gaps."
        )

        if matching:
            doc.add_paragraph("Highlights of my fit include:")
            for s in matching:
                doc.add_paragraph(str(s), style="List Bullet")

        if missing:
            doc.add_paragraph(
                "I’m also proactively upskilling in the following areas to meet the role’s requirements:"
            )
            for s in missing:
                doc.add_paragraph(str(s), style="List Bullet")

        doc.add_paragraph(
            "I would welcome the opportunity to discuss how my experience can help your team deliver impact. "
            "Thank you for your time and consideration."
        )
        doc.add_paragraph("Sincerely,")
        doc.add_paragraph("[Your Name]")

        filename = _safe_filename(f"cover_letter_{analysis_id}.docx")

    buf = io.BytesIO()
    doc.save(buf)
    return DocxExportResult(
        filename=filename,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        data=buf.getvalue(),
    )





